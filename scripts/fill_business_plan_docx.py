#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Заполнение шаблона «Бизнес план.docx» данными проекта МаБибип (350 000 ₽)."""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "Бизнес план.docx"
OUTPUT = ROOT / "docs" / "BUSINESS-PLAN-MABIBIP-350K.docx"

# --- Финансовая модель (согласована с BUSINESS-PLAN-MABIBIP-350K-FILLED.md) ---
BASE_REVENUE = 300_000
COEFFS = [0.12, 0.15, 0.18, 0.22, 0.28, 0.35, 0.43, 0.52, 0.64, 0.76, 0.88, 1.00]
MONTHS = [
    "1 месяц",
    "2 месяц",
    "3 месяц",
    "4 месяц",
    "5 месяц",
    "6 месяц",
    "7 месяц",
    "8 месяц",
    "9 месяц",
    "10 месяц",
    "11 месяц",
    "12 месяц",
]
REVENUES = [int(BASE_REVENUE * c) for c in COEFFS]
FIXED_MONTHLY = 8_000
ONE_TIME = 350_000
VAR_RATE = 0.05

EXPENSES: list[int] = []
for i, rev in enumerate(REVENUES):
    var = int(rev * VAR_RATE)
    EXPENSES.append(ONE_TIME + FIXED_MONTHLY + var if i == 0 else FIXED_MONTHLY + var)

TAXES = [int(r * 0.06) for r in REVENUES]
NET = [REVENUES[i] - EXPENSES[i] - TAXES[i] for i in range(12)]

# --- Смета 350 000 ₽ ---
BUDGET_LINES = [
    (
        "1.1",
        "Ноутбук Apple MacBook (Apple Silicon, RAM 16 GB, SSD 512 GB) — разработка и сопровождение платформы",
        "1",
        "120 000",
        "120 000",
        "re:Store / DNS / Biggeek",
    ),
    (
        "1.2",
        "Смартфон Android — тестирование мобильной версии, связь с партнёрами",
        "1",
        "22 000",
        "22 000",
        "DNS / Ozon / М.Видео",
    ),
    (
        "1.3",
        "Внешний SSD 1 TB — резервное копирование и перенос данных проекта",
        "1",
        "8 000",
        "8 000",
        "DNS / Ozon",
    ),
    (
        "1.4",
        "Чехол и защитные аксессуары для ноутбука",
        "1",
        "5 000",
        "5 000",
        "re:Store / Ozon",
    ),
    (
        "1.5",
        "Компьютерная мышь (эргономичная, беспроводная)",
        "1",
        "3 000",
        "3 000",
        "DNS / Ozon",
    ),
    (
        "2.1",
        "Кресло компьютерное с ортопедической поддержкой (рабочее место)",
        "1",
        "15 000",
        "15 000",
        "Hoff / Ozon / Ikea",
    ),
    (
        "2.2",
        "МФУ / принтер (лазерный) — печать договоров, визиток, буклетов",
        "1",
        "20 000",
        "20 000",
        "DNS / М.Видео",
    ),
    (
        "3.1",
        "Аренда VPS / облачного сервера (12 мес.) — размещение mabibip.ru",
        "1",
        "20 000",
        "20 000",
        "Beget / Selectel / Timeweb",
    ),
    (
        "3.2",
        "Подписки на ПО для разработки (IDE, Git, дизайн, тестирование), 12 мес.",
        "1",
        "15 000",
        "15 000",
        "JetBrains / GitHub / Figma",
    ),
    (
        "4.1",
        "Абонентская плата за домашний интернет / Wi‑Fi (12 мес.) — удалённая работа",
        "1",
        "10 000",
        "10 000",
        "Ростелеком / Эр‑Телеком",
    ),
    (
        "5.1",
        "Комплексное продвижение: интернет‑реклама (VK, Яндекс), соцсети; полиграфия "
        "(буклеты, листовки, визитки); QR‑ и POS‑материалы для партнёров",
        "1",
        "100 000",
        "100 000",
        "VK Ads / Яндекс Директ / типография",
    ),
    (
        "6.1",
        "Транспортные расходы — встречи с мастерами, СТО и партнёрами (г. Владикавказ)",
        "1",
        "7 000",
        "7 000",
        "—",
    ),
    (
        "6.2",
        "Домен mabibip.ru + SSL‑сертификат, 12 мес.",
        "1",
        "5 000",
        "5 000",
        "REG.RU / nic.ru",
    ),
]

BUDGET_SUMMARY = [
    ("Основные средства", "185 000", "MacBook, смартфон, SSD, кресло, принтер"),
    ("Материально‑производственные запасы", "8 000", "Чехол, компьютерная мышь"),
    ("ПО, хостинг и подписки для разработки", "35 000", "VPS и лицензии ПО (10% от сметы)"),
    ("Услуги связи (интернет / Wi‑Fi)", "10 000", "12 месяцев"),
    ("Продвижение и маркетинг", "100 000", "Реклама, полиграфия, материалы для партнёров"),
    ("Транспортные расходы", "7 000", "Разъезды по партнёрам"),
    ("Прочее (домен, SSL)", "5 000", "REG.RU / nic.ru"),
]

MONTHLY_EXPENSE_ROWS = [
    ("Аренда помещения, оборудования", "0", "0", "0"),
    ("Коммунальные услуги", "0", "0", "0"),
    ("Налоги на имущество", "0", "0", "0"),
    ("Связь (мобильная связь)", "1 500", "18 000", "18 000"),
    ("Транспортные расходы", "500", "6 000", "6 000"),
    ("ГСМ", "0", "0", "0"),
    ("Канцелярские расходы", "500", "6 000", "6 000"),
    ("Ремонт и обслуживание", "0", "0", "0"),
    ("Реклама (поддерживающая, из выручки)", "3 000", "36 000", "36 000"),
    ("Разрешительная документация", "0", "0", "0"),
    ("Банковские услуги и комиссии", "500", "6 000", "6 000"),
    ("Оплата товаров и услуг (SMS, продление сервисов)", "2 000", "24 000", "24 000"),
    ("Заработная плата персонала", "0", "0", "0"),
    ("Прочие расходы", "0", "0", "0"),
]

DIRECT_COSTS = [
    # 4.2.2 — выручка по тарифам (год 1 / год 2), согласовано с помесячной моделью
    ("Подписка B2C «Активный» (199 ₽/мес.)", "220 329", "275 411"),
    ("Подписка B2C «Профи» (399 ₽/мес.)", "176 706", "220 882"),
    ("Подписка «Мастер» (990 ₽/мес.)", "274 027", "342 533"),
    ("Подписка СТО «Бизнес» (2 990 ₽/мес.)", "248 285", "310 356"),
    ("Подписка СТО «Корпоративный» (5 990 ₽/мес.)", "132 640", "165 800"),
    ("Подписка магазина «Витрина» (1 490 ₽/мес.)", "164 969", "206 211"),
    ("Подписка магазина «Магазин» (3 990 ₽/мес.)", "220 882", "276 102"),
    ("Подписка магазина «Сеть» (7 990 ₽/мес.)", "221 159", "276 448"),
]


def _set_cell(row, idx: int, text: str) -> None:
    if idx < len(row.cells):
        row.cells[idx].text = text


def _replace_paragraph(doc: Document, idx: int, text: str) -> None:
    if 0 <= idx < len(doc.paragraphs):
        doc.paragraphs[idx].text = text


def _normalize_region_text(text: str) -> str:
    """Заменить остатки шаблона Ярославской области на РСО‑Алania / Владикавказ."""
    if not text or "Ярослав" not in text:
        return text
    replacements = [
        (
            "Ярославской области/города Ярославля",
            "Республики Северная Осетия — Алания/г. Владикавказ",
        ),
        (
            "муниципального  района Ярославской области/города Ярославля",
            "республики Северная Осетия — Алания/г. Владикавказ",
        ),
        (
            "муниципальные  образования  Ярославской области",
            "муниципальные образования Республики Северная Осетия — Алания (г. Владикавказ)",
        ),
        ("муниципальные образования Ярославской области", "муниципальные образования Республики Северная Осетия — Алания (г. Владикавказ)"),
        ("Ярославской области", "Республики Северная Осетия — Алания"),
        ("города Ярославля", "г. Владикавказ"),
        ("Мэрия города Ярославля", "Республика Северная Осетия — Алания, г. Владикавказ"),
    ]
    result = text
    for old, new in replacements:
        result = result.replace(old, new)
    result = re.sub(r"Ярослав\w*", "Владикавказ", result)
    return result


def _apply_region_replacements(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        normalized = _normalize_region_text(paragraph.text)
        if normalized != paragraph.text:
            paragraph.text = normalized
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    normalized = _normalize_region_text(paragraph.text)
                    if normalized != paragraph.text:
                        paragraph.text = normalized
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in header_footer.paragraphs:
                normalized = _normalize_region_text(paragraph.text)
                if normalized != paragraph.text:
                    paragraph.text = normalized


def _patch_docx_metadata(docx_path: Path) -> None:
    """Убрать «Мэрия города Ярославля» из docProps/app.xml."""
    with zipfile.ZipFile(docx_path, "r") as zin:
        parts = {name: zin.read(name) for name in zin.namelist()}
    app_xml = parts.get("docProps/app.xml")
    if not app_xml:
        return
    text = app_xml.decode("utf-8")
    patched = _normalize_region_text(text)
    if patched == text:
        return
    parts["docProps/app.xml"] = patched.encode("utf-8")
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    tmp.replace(docx_path)


def fill_document(doc: Document) -> None:
    # Шапка
    _replace_paragraph(doc, 2, "БИЗНЕС‑ПЛАН")
    _replace_paragraph(doc,
        3,
        "Социальный контракт на осуществление индивидуальной предпринимательской деятельности. "
        "Заявитель: Скяев Алихан Артурович.",
    )
    _replace_paragraph(doc,
        5,
        "Наименование проекта: «МаБибип» — цифровой сервис для автовладельцев и автобизнеса "
        "(https://mabibip.ru): каталог мастеров и СТО, онлайн‑запись, доска объявлений.",
    )

    # Раздел 1 — по структуре шаблона (п. 1.1–1.13)
    _replace_paragraph(doc, 9, "Скяев Алихан Артурович")
    _replace_paragraph(doc,
        11,
        "Телефон: +7 (918) 702‑09‑87. E-mail: alihanskaev@gmail.com.",
    )
    _replace_paragraph(doc,
        12,
        "1.3. Адрес по прописке: Республика Северная Осетия — Алания, г. Владикавказ, "
        "ул. Морских Пехотинцев, д. 5, кв. 100.",
    )
    _replace_paragraph(doc,
        13,
        "1.4. Фактический адрес проживания: совпадает с адресом по прописке.",
    )
    _replace_paragraph(doc,
        14,
        "1.5. Резюме: дата рождения — 21.02.2002; семейное положение — холост; состав семьи — "
        "1 человек (заявитель). Образование: среднее общее (СКСВУ, г. Владикавказ, 2020 г.); "
        "высшее — обучается на 4 курсе бакалавриата (прикладная информатика / IT). "
        "Статус: зарегистрирован в органах службы занятости как безработный.",
    )
    _replace_paragraph(doc,
        16,
        "По найму не работал. С 2024 г. — самостоятельная разработка и сопровождение "
        "веб‑платформы mabibip.ru (каталог мастеров и СТО, объявления, онлайн‑запись, чаты).",
    )
    _replace_paragraph(doc,
        18,
        "Разработка и сопровождение веб‑сервисов (Python/Django), администрирование серверов "
        "(Linux/VPS), наполнение каталога, интернет‑маркетинг, переговоры с партнёрами в автосфере.",
    )
    _replace_paragraph(doc,
        23,
        "Разработка, запуск и развитие платформы «МаБибип». "
        "ОКВЭД 62.01 «Разработка компьютерного ПО» (основной); дополнительно: 63.12 «Web‑порталы», "
        "62.09 «Прочая деятельность в области IT». Форма: самозанятый (НПД), учёт через «Мой налог».",
    )
    _replace_paragraph(doc,
        25,
        "Место деятельности: дистанционно по месту проживания (г. Владикавказ, "
        "ул. Морских Пехотинцев, д. 5, кв. 100) — домашнее рабочее место, собственность / без аренды.",
    )
    _replace_paragraph(doc,
        27,
        "Не требуется: деятельность осуществляется дистанционно (IT‑сервис), "
        "отдельное производственное помещение и дооборудование не предусмотрены.",
    )
    _replace_paragraph(doc,
        30,
        "Кадровый состав: 0 наёмных работников (самозанятость). "
        "1.12. Имеется: действующий сайт mabibip.ru; ноутбук (личный); опыт разработки; "
        "база партнёров‑мастеров для подключения. По смете соцконтракта приобретаются: MacBook, "
        "смартфон, SSD, рабочее место, сервер, ПО, рекламные материалы.",
    )
    _replace_paragraph(doc,
        32,
        "Опыт разработки IT‑сервисов — 2 года (2024–2026); опыт ведения данного вида "
        "предпринимательской деятельности как самозанятого — с момента регистрации НПД. "
        "Срок реализации проекта по соцконтракту: 12 месяцев.",
    )

    # Раздел 2
    _replace_paragraph(doc,
        35,
        "2.1. Цель: запустить и развить в РСО‑Алания платформу «МаБибип» — единый сервис "
        "для поиска мастеров и СТО, объявлений, онлайн‑записи и связи с исполнителями. "
        "Задачи: (1) доработка продукта; (2) подключение мастеров, СТО и магазинов; "
        "(3) продвижение; (4) внедрение подписок; (5) модерация и антифрод.",
    )
    _replace_paragraph(doc,
        37,
        "2.2. Обоснование соцпомощи: заявитель зарегистрирован как безработный, "
        "имеет IT‑компетенции и действующий прототип сервиса, но не располагает средствами "
        "на технику, сервер, рекламу и рабочее место, необходимые для выхода на "
        "самообеспечение. Средства соцконтракта (350 000 ₽) направляются на стартовые "
        "вложения; доход — от подписок после выхода на операционную окупаемость (7–8 мес.).",
    )
    _replace_paragraph(doc,
        39,
        "2.3. Актуальность для РСО‑Алания и г. Владикавказа: цифровизация рынка автосервисов "
        "повышает доступность услуг, создаёт занятость для мастеров и СТО, снижает «серый» рынок. "
        "Направление: информационные технологии, онлайн‑сервис для автомобильной тематики.",
    )
    _replace_paragraph(doc,
        29,
        "1.11. Кадровый состав: 0 наёмных работников (самозанятость, без найма).",
    )
    _replace_paragraph(doc,
        40,
        "2.4. Конкуренты на рынке: Авито, 2ГИС, Яндекс Карты, локальные чаты. "
        "Специализированной платформы с онлайн‑записью и кабинетом для бизнеса в регионе нет.",
    )
    _replace_paragraph(doc,
        41,
        "График работы: 6 дней в неделю, 36–48 часов. Подготовительный этап — 1–3 месяца. "
        "Окупаемость — 7–8 месяцев. Наёмный персонал не планируется.",
    )

    # Раздел 3
    _replace_paragraph(doc,
        45,
        "3.1. Целевая аудитория: водители 18–65 лет; исполнители — мастера, СТО, магазины. "
        "География: РСО‑Алания, г. Владикавказ. "
        "Договоры о намерении: планируется получение 2–5 писем от мастеров и СТО "
        "(партнёры из списка подключения; приложить к заявке).",
    )
    _replace_paragraph(doc,
        47,
        "3.3. Продвижение: VK Реклама, Яндекс Директ, SEO, QR‑материалы у партнёров, "
        "буклеты и визитки. Монетизация — подписки: B2C («Активный» 199 ₽, «Профи» 399 ₽); "
        "мастера («Мастер» 990 ₽); СТО («Бизнес» 2 990 ₽, «Корпоративный» 5 990 ₽); "
        "магазины («Витрина» 1 490 ₽, «Магазин» 3 990 ₽, «Сеть» 7 990 ₽).",
    )

    # Таблица 0 — трудовая деятельность (п. 1.6)
    t0 = doc.tables[0]
    work_rows = [
        ("2024 — н.в.", "Разработка и запуск платформы «МаБибип» (самостоятельно)"),
        ("2020 — 2024", "Обучение в ВУЗе, практика в области IT"),
        ("Не применимо", "Иная официальная занятость по найму отсутствует"),
    ]
    for i, (period, role) in enumerate(work_rows, start=1):
        if i < len(t0.rows):
            _set_cell(t0.rows[i], 0, period)
            _set_cell(t0.rows[i], 1, role)
    for i in range(len(work_rows) + 1, len(t0.rows)):
        _set_cell(t0.rows[i], 0, "—")
        _set_cell(t0.rows[i], 1, "Не применимо")

    # Таблица 1 — дополнительные знания (п. 1.7)
    t1 = doc.tables[1]
    skill_rows = [
        ("Веб‑разработка (Python/Django)", "Создание и сопровождение mabibip.ru"),
        ("Администрирование серверов", "Linux, VPS, Docker, резервное копирование"),
        ("Интернет‑маркетинг", "VK Ads, Яндекс Директ, SEO, работа с партнёрами"),
        ("Самостоятельное обучение", "Право ПДн, учёт НПД, финансовое планирование"),
    ]
    for i, (course, skills) in enumerate(skill_rows, start=1):
        if i < len(t1.rows):
            _set_cell(t1.rows[i], 0, course)
            _set_cell(t1.rows[i], 1, skills)

    # Таблица 2 — свод 4.1.1
    t2 = doc.tables[2]
    while len(t2.rows) < len(BUDGET_SUMMARY) + 3:
        t2.add_row()
    for i, (name, amount, note) in enumerate(BUDGET_SUMMARY, start=2):
        row = t2.rows[i]
        num = str(i - 1)
        _set_cell(row, 0, num)
        _set_cell(row, 1, f"{num}. {name}" if not name.startswith(num) else name)
        _set_cell(row, 2, amount)
        _set_cell(row, 3, amount)
        _set_cell(row, 4, "0")
        _set_cell(row, 5, "0")
    total_row = len(BUDGET_SUMMARY) + 2
    tr = t2.rows[total_row]
    _set_cell(tr, 0, "")
    _set_cell(tr, 1, "ИТОГО")
    _set_cell(tr, 2, "350 000")
    _set_cell(tr, 3, "350 000")
    _set_cell(tr, 4, "0")
    _set_cell(tr, 5, "0")

    # Таблица 3 — детализация 4.1.2
    t3 = doc.tables[3]
    # Очистить пример швейной машины
    while len(t3.rows) > 2:
        tbl = t3._tbl
        tbl.remove(t3.rows[-1]._tr)
    for line in BUDGET_LINES:
        row = t3.add_row()
        _set_cell(row, 0, line[0])
        _set_cell(row, 1, line[1])
        _set_cell(row, 2, line[3].replace(" ", ""))
        _set_cell(row, 3, line[2])
        _set_cell(row, 4, line[4].replace(" ", ""))
    total = t3.add_row()
    _set_cell(total, 1, "ИТОГО")
    _set_cell(total, 4, "350000")

    # Таблица 4 — выручка 4.2.1 (помесячный прогноз)
    t4 = doc.tables[4]
    while len(t4.rows) > 2:
        t4._tbl.remove(t4.rows[-1]._tr)
    note_row = t4.rows[1]
    _set_cell(note_row, 1, "Помесячный прогноз (12 мес.)")
    _set_cell(note_row, 2, "коэфф.")
    _set_cell(note_row, 3, "выручка")
    _set_cell(note_row, 4, "расходы")
    _set_cell(note_row, 5, "прибыль до налога")
    for i, month in enumerate(MONTHS):
        row = t4.add_row()
        _set_cell(row, 0, str(i + 1))
        _set_cell(row, 1, month)
        _set_cell(row, 2, f"{COEFFS[i]:.2f}".replace(".", ","))
        _set_cell(row, 3, str(REVENUES[i]))
        _set_cell(row, 4, str(EXPENSES[i]))
        profit = REVENUES[i] - EXPENSES[i]
        _set_cell(row, 5, str(profit))
    sum_row = t4.add_row()
    _set_cell(sum_row, 1, "ИТОГО за 12 месяцев")
    _set_cell(sum_row, 3, str(sum(REVENUES)))
    _set_cell(sum_row, 4, str(sum(EXPENSES)))
    _set_cell(sum_row, 5, str(sum(REVENUES) - sum(EXPENSES)))

    # Таблица 5 — выручка по тарифам 4.2.2
    t5 = doc.tables[5]
    while len(t5.rows) > 2:
        t5._tbl.remove(t5.rows[-1]._tr)
    for i, (name, y1, y2) in enumerate(DIRECT_COSTS, start=1):
        row = t5.add_row()
        _set_cell(row, 0, str(i))
        _set_cell(row, 1, name)
        _set_cell(row, 2, y1.replace(" ", ""))
        _set_cell(row, 3, y2.replace(" ", ""))
    total5 = t5.add_row()
    _set_cell(total5, 0, "9")
    _set_cell(total5, 1, "ИТОГО выручка")
    _set_cell(total5, 2, str(sum(REVENUES)))
    _set_cell(total5, 3, str(int(sum(REVENUES) * 1.25)))

    # Таблица 6 — ежемесячные расходы 4.2.3
    t6 = doc.tables[6]
    for i, (name, monthly, y1, y2) in enumerate(MONTHLY_EXPENSE_ROWS, start=2):
        if i >= len(t6.rows):
            t6.add_row()
        row = t6.rows[i]
        _set_cell(row, 0, str(i - 1))
        _set_cell(row, 1, name)
        _set_cell(row, 2, monthly)
        _set_cell(row, 3, y1)
        _set_cell(row, 4, y2)
    if len(t6.rows) > len(MONTHLY_EXPENSE_ROWS) + 2:
        total_i = len(MONTHLY_EXPENSE_ROWS) + 2
        if total_i < len(t6.rows):
            tr6 = t6.rows[total_i]
            _set_cell(tr6, 0, "")
            _set_cell(tr6, 1, "ИТОГО")
            _set_cell(tr6, 2, "8 000")
            _set_cell(tr6, 3, "96 000")
            _set_cell(tr6, 4, "96 000")
    if len(t6.rows) > 17:
        extra = t6.rows[17]
        _set_cell(extra, 0, "")
        _set_cell(extra, 1, "")
        _set_cell(extra, 2, "")
        _set_cell(extra, 3, "")
        _set_cell(extra, 4, "")

    # Таблица 7 — налоги и прибыль 4.2.4
    t7 = doc.tables[7]
    avg_tax = sum(TAXES) // 12
    year_tax = sum(TAXES)
    year2_tax = int(sum(REVENUES) * 1.25 * 0.06)
    profit_y1 = sum(REVENUES) - sum(EXPENSES)
    profit_y2 = int(profit_y1 * 1.25)
    net_y1 = profit_y1 - year_tax
    net_y2 = int(net_y1 * 1.25)
    tax_rows = [
        ("НПД с доходов от физических лиц (6%)", str(avg_tax), str(year_tax), str(year2_tax)),
        ("НПД с доходов от юридических лиц (4%)", "0", "0", "0"),
        ("Прибыль до налогообложения", str(profit_y1 // 12), str(profit_y1), str(profit_y2)),
        ("Чистая прибыль (после НПД)", str(net_y1 // 12), str(net_y1), str(net_y2)),
    ]
    for i, (name, m, y1, y2) in enumerate(tax_rows, start=2):
        if i >= len(t7.rows):
            t7.add_row()
        row = t7.rows[i]
        _set_cell(row, 0, str(i - 1))
        _set_cell(row, 1, name)
        _set_cell(row, 2, m)
        _set_cell(row, 3, y1)
        _set_cell(row, 4, y2)
    _replace_paragraph(doc, 64, "Скяев А. А. _____________________ «___» __________ 2026 г.")

    _apply_region_replacements(doc)


def main() -> None:
    if not TEMPLATE.is_file():
        raise SystemExit(f"Шаблон не найден: {TEMPLATE}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))
    fill_document(doc)
    doc.save(str(OUTPUT))
    _patch_docx_metadata(OUTPUT)
    print(f"Готово: {OUTPUT}")


if __name__ == "__main__":
    main()
