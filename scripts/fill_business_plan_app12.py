#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Шаблон и заполнение бизнес-плана по форме Приложение 12 (Постановление №434, ред. 09.09.2025).
Проект: МаБибип, смета 350 000 ₽.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.bp_app12_form_text import (  # noqa: E402
    APPENDIX_HEADER_LINES,
    FINANCE_TABLE_REFS,
    MARKETING_TABLE_REFS,
    MARKET_ITEMS,
    MARKETING_INTRO,
    SECTION1_ITEMS,
    SECTION2_ITEMS,
)
from scripts.bp_app12_finance import (  # noqa: E402
    ADS_SUPPORT_ROW,
    CAPEX_ROW,
    COEFFS,
    EXPENSES,
    INTERNET_ROW,
    LIMIT_PROMO,
    LIMIT_RENT,
    LIMIT_SW,
    MONTH_LABELS,
    NET_PROFIT,
    PROFIT_BEFORE_TAX,
    RENT_ROW,
    REVENUES,
    TABLE1_LINES,
    TABLE1_SUMS,
    TABLE1_TOTAL,
    TABLE2_LINES,
    TABLE3_LINES,
    TABLE3_TOTAL_DIRECT,
    TABLE3_TOTAL_REV,
    TABLE4_LINES,
    TABLE4_TOTAL,
    TABLE6_ROWS,
    TABLE8_RISKS,
    TAXES,
    VAR_COSTS,
    YEAR_COGS,
    YEAR_EXPENSES,
    YEAR_NET,
    YEAR_REVENUE,
    YEAR_TAX,
)

TEMPLATE_PATH = ROOT / "docs" / "templates" / "Приложение-12-шаблон.docx"
OUTPUT_PATH = ROOT / "docs" / "BUSINESS-PLAN-MABIBIP-350K-APP12.docx"
OUTPUT_REGENERATED = ROOT / "docs" / "BUSINESS-PLAN-MABIBIP-350K-APP12-regenerated.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
FONT_SIZE_TITLE = Pt(14)
LINE_SPACING = 1.15
MARGIN_CM = 2.0


def _fmt(n: int) -> str:
    return f"{n:,}".replace(",", " ")


def _set_run_font(
    run,
    *,
    bold: bool = False,
    size=None,
    underline: bool = False,
) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = size or FONT_SIZE
    run.bold = bold
    if underline:
        run.underline = WD_UNDERLINE.SINGLE


def _format_body_paragraph(p, *, align=None) -> None:
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.first_line_indent = Cm(0)


def _add_para(doc: Document, text: str, *, bold: bool = False, align=None, size=None) -> None:
    p = doc.add_paragraph()
    _format_body_paragraph(p, align=align)
    run = p.add_run(text)
    _set_run_font(run, bold=bold, size=size)


def _add_appendix_header(doc: Document) -> None:
    """Шапка «Приложение 12» — как в официальной форме (справа вверху первой страницы)."""
    for line in APPENDIX_HEADER_LINES:
        _add_para(doc, line, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()


def _add_form_field(
    doc: Document,
    num: str,
    label: str,
    answer: str,
    *,
    sub: bool = False,
) -> None:
    """Формулировка пункта обычным текстом, ответ — с подчёркиванием."""
    p = doc.add_paragraph()
    _format_body_paragraph(p)
    if sub:
        p.paragraph_format.left_indent = Cm(0.75)
    prefix = f"{num}) {label}: "
    r1 = p.add_run(prefix)
    _set_run_font(r1)
    r2 = p.add_run(answer)
    _set_run_font(r2, underline=True)


def _add_table_ref(doc: Document, num: str, label: str) -> None:
    p = doc.add_paragraph()
    _format_body_paragraph(p)
    run = p.add_run(f"{num}) {label}")
    _set_run_font(run)


def _add_section_heading(doc: Document, text: str) -> None:
    _add_para(doc, text, bold=True)


def _add_form_fields(doc: Document, items: list[tuple[str, str, str]]) -> None:
    for num, label, answer in items:
        _add_form_field(doc, num, label, answer, sub=num in ("а", "б"))


def _style_cell(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    _format_body_paragraph(p)
    run = p.add_run(str(text))
    _set_run_font(run, bold=bold)


def _add_table(doc: Document, rows: int, cols: int):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = "Table Grid"
    return tbl


def _setup_document(doc: Document) -> None:
    m = Cm(MARGIN_CM)
    for section in doc.sections:
        section.top_margin = m
        section.bottom_margin = m
        section.left_margin = m
        section.right_margin = m
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING


def build_template(doc: Document | None = None) -> Document:
    """Пустая форма Приложение 12 со всеми таблицами и полями-заполнителями."""
    doc = doc or Document()
    _setup_document(doc)

    _add_para(doc, "БИЗНЕС-ПЛАН", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=FONT_SIZE_TITLE)
    _add_para(doc, "(форма согласно Приложению 12 к Постановлению Правительства РСО-Алания №434)")
    doc.add_paragraph()

    _add_para(doc, "1. Информация о заявителе:", bold=True)
    tpl1 = [
        "1) фамилия, имя, отчество (при наличии): _________________________________",
        "2) дата рождения: _________________________________",
        "3) место жительства: _________________________________",
        "4) е-mail, телефон: _________________________________",
        "5) состав семьи (количество человек): _________________________________",
        "6) образование (специальность)...: _________________________________",
        "7) дополнительные знания, умения, навыки...: _________________________________",
        "8) потребность в обучении/повышении квалификации...: _________________________________",
    ]
    for line in tpl1:
        _add_para(doc, line)
    doc.add_paragraph()

    _add_para(doc, "2. Сведения о проекте", bold=True)
    for i in range(1, 13):
        _add_para(doc, f"{i}) _________________________________________________")
    doc.add_paragraph()

    _add_para(doc, "таблица 1", bold=True)
    t1 = _add_table(doc, 2, 7)
    hdr1 = ["№", "Наименование", "Кол-во", "Цена, руб.", "Сумма, руб.", "Поставщик", "Примечание"]
    for j, h in enumerate(hdr1):
        _style_cell(t1.rows[0].cells[j], h, bold=True)
    _style_cell(t1.rows[1].cells[0], "…")
    doc.add_paragraph()

    _add_para(doc, "таблица 2", bold=True)
    t2 = _add_table(doc, 2, 7)
    hdr2 = ["№", "Наименование", "Назначение", "Кол-во", "Эконом", "Стандарт", "Премиум"]
    for j, h in enumerate(hdr2):
        _style_cell(t2.rows[0].cells[j], h, bold=True)
    doc.add_paragraph()

    _add_para(doc, "Анализ рынка:", bold=True)
    for i in range(1, 5):
        _add_para(doc, f"{i}) _________________________________________________")
    doc.add_paragraph()

    _add_para(doc, "3. Маркетинговый план", bold=True)
    _add_para(doc, "таблица 3", bold=True)
    t3 = _add_table(doc, 2, 7)
    hdr3 = ["Услуга/товар", "Ед.", "Кол-во/мес.", "Цена", "Выручка", "Прямые %", "Прямые руб."]
    for j, h in enumerate(hdr3):
        _style_cell(t3.rows[0].cells[j], h, bold=True)
    doc.add_paragraph()

    _add_para(doc, "таблица 4", bold=True)
    t4 = _add_table(doc, 2, 2)
    _style_cell(t4.rows[0].cells[0], "Статья расходов", bold=True)
    _style_cell(t4.rows[0].cells[1], "Сумма в месяц, руб.", bold=True)
    doc.add_paragraph()

    _add_para(doc, "4. Финансовый план", bold=True)
    _add_para(doc, "таблица 5", bold=True)
    t5 = _add_table(doc, 2, 14)
    for j in range(14):
        _style_cell(t5.rows[0].cells[j], f"М{j}" if j else "Показатель", bold=True)
    doc.add_paragraph()

    _add_para(doc, "таблица 6", bold=True)
    t6 = _add_table(doc, 2, 4)
    hdr6 = ["Показатель", "Ед.", "Среднемесячно", "За год"]
    for j, h in enumerate(hdr6):
        _style_cell(t6.rows[0].cells[j], h, bold=True)
    doc.add_paragraph()

    _add_para(doc, "таблица 7", bold=True)
    t7 = _add_table(doc, 2, 3)
    _style_cell(t7.rows[0].cells[0], "Источник", bold=True)
    _style_cell(t7.rows[0].cells[1], "Сумма, руб.", bold=True)
    _style_cell(t7.rows[0].cells[2], "Доля, %", bold=True)
    doc.add_paragraph()

    _add_para(doc, "таблица 8", bold=True)
    t8 = _add_table(doc, 2, 2)
    _style_cell(t8.rows[0].cells[0], "Риск", bold=True)
    _style_cell(t8.rows[0].cells[1], "Меры снижения", bold=True)

    return doc


def fill_document(doc: Document) -> Document:
    """Заполняет документ данными МаБибип (пересобирает содержимое)."""
    doc = Document()
    _setup_document(doc)

    _add_appendix_header(doc)

    _add_para(doc, "БИЗНЕС-ПЛАН", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=FONT_SIZE_TITLE)
    doc.add_paragraph()

    # --- Раздел 1 ---
    _add_section_heading(doc, "1. Информация о заявителе:")
    _add_form_fields(doc, SECTION1_ITEMS)
    doc.add_paragraph()

    # --- Раздел 2 проект ---
    _add_section_heading(doc, "2. Описание проекта:")
    _add_form_fields(doc, SECTION2_ITEMS)
    doc.add_paragraph()

    # --- Таблица 1 ---
    _add_para(doc, "таблица 1", bold=True)
    _add_para(
        doc,
        "Смета составлена с соблюдением лимитов Правил: п. 5 — не более 5% (17 500 руб.); "
        "п. 4 — не более 10% (35 000 руб.); п. 3 — не более 15% (52 500 руб.). "
        "Расширенное продвижение после запуска — за счёт выручки (таблица 4, 3 000 руб./мес.).",
    )
    rows_t1 = 1 + len(TABLE1_LINES) + 2
    t1 = _add_table(doc, rows_t1, 7)
    hdr1 = ["№", "Наименование", "Кол-во", "Цена, руб.", "Сумма, руб.", "Поставщик", "Примечание"]
    for j, h in enumerate(hdr1):
        _style_cell(t1.rows[0].cells[j], h, bold=True)
    r = 1
    for line in TABLE1_LINES:
        code, name, qty, price, total, supplier, is_hdr = line
        if is_hdr:
            _style_cell(t1.rows[r].cells[0], code, bold=True)
            _style_cell(t1.rows[r].cells[1], name, bold=True)
        else:
            _style_cell(t1.rows[r].cells[0], code)
            _style_cell(t1.rows[r].cells[1], name)
            _style_cell(t1.rows[r].cells[2], qty)
            _style_cell(t1.rows[r].cells[3], price)
            _style_cell(t1.rows[r].cells[4], total)
            _style_cell(t1.rows[r].cells[5], supplier)
        r += 1
    _style_cell(t1.rows[r].cells[0], "ИТОГО", bold=True)
    _style_cell(t1.rows[r].cells[4], _fmt(TABLE1_TOTAL), bold=True)
    doc.add_paragraph()

    # --- Таблица 2 ---
    _add_para(doc, "13) анализ цен на рынке (таблица 2)", bold=True)
    t2 = _add_table(doc, 1 + len(TABLE2_LINES), 7)
    hdr2 = ["№", "Наименование", "Назначение", "Кол-во", "Эконом", "Стандарт", "Премиум"]
    for j, h in enumerate(hdr2):
        _style_cell(t2.rows[0].cells[j], h, bold=True)
    for i, row in enumerate(TABLE2_LINES, start=1):
        for j, val in enumerate(row):
            _style_cell(t2.rows[i].cells[j], val)
    doc.add_paragraph()

    # --- Рынок ---
    _add_section_heading(doc, "Анализ рынка и конкурентов:")
    _add_form_fields(doc, MARKET_ITEMS)
    doc.add_paragraph()

    # --- Раздел 3 ---
    _add_section_heading(doc, "3. Маркетинговый план:")
    _add_form_fields(doc, MARKETING_INTRO)
    _add_table_ref(doc, *MARKETING_TABLE_REFS[0])
    t3 = _add_table(doc, 1 + len(TABLE3_LINES) + 1, 7)
    hdr3 = ["Услуга", "Ед.", "Подписч./мес.", "Цена, руб.", "Выручка", "Прямые %", "Прямые руб."]
    for j, h in enumerate(hdr3):
        _style_cell(t3.rows[0].cells[j], h, bold=True)
    for i, row in enumerate(TABLE3_LINES, start=1):
        for j, val in enumerate(row):
            _style_cell(t3.rows[i].cells[j], val)
    ri = len(TABLE3_LINES) + 1
    _style_cell(t3.rows[ri].cells[0], "ИТОГО", bold=True)
    _style_cell(t3.rows[ri].cells[4], _fmt(TABLE3_TOTAL_REV), bold=True)
    _style_cell(t3.rows[ri].cells[6], _fmt(TABLE3_TOTAL_DIRECT), bold=True)
    doc.add_paragraph()

    _add_table_ref(doc, *MARKETING_TABLE_REFS[1])
    t4 = _add_table(doc, 1 + len(TABLE4_LINES) + 1, 2)
    _style_cell(t4.rows[0].cells[0], "Статья расходов", bold=True)
    _style_cell(t4.rows[0].cells[1], "Сумма в месяц, руб.", bold=True)
    for i, (name, amt) in enumerate(TABLE4_LINES, start=1):
        _style_cell(t4.rows[i].cells[0], name)
        _style_cell(t4.rows[i].cells[1], amt)
    ti = len(TABLE4_LINES) + 1
    _style_cell(t4.rows[ti].cells[0], "ИТОГО", bold=True)
    _style_cell(t4.rows[ti].cells[1], _fmt(TABLE4_TOTAL), bold=True)
    doc.add_paragraph()

    # --- Раздел 4 / Таблица 5 ---
    _add_section_heading(doc, "4. Финансовый план:")
    _add_table_ref(doc, *FINANCE_TABLE_REFS[0])

    table5_rows = [
        ("1", "Наименование месяца", MONTH_LABELS),
        ("2", "Коэффициент сезонности", [str(c).replace(".", ",") for c in COEFFS]),
        ("3", "Выручка", [_fmt(r) for r in REVENUES]),
        ("4", "Расходы всего", [_fmt(e) for e in EXPENSES]),
        ("4.1", "в т.ч. аренда", [_fmt(x) for x in RENT_ROW]),
        ("4.2", "в т.ч. основные средства (смета)", [_fmt(x) for x in CAPEX_ROW]),
        ("4.3", "в т.ч. переменные (комиссии ~5%)", [_fmt(x) for x in VAR_COSTS]),
        ("4.4", "в т.ч. интернет", [_fmt(x) for x in INTERNET_ROW]),
        ("4.5", "в т.ч. реклама поддерживающая", [_fmt(x) for x in ADS_SUPPORT_ROW]),
        ("5", "Прибыль до налогообложения", [_fmt(p) for p in PROFIT_BEFORE_TAX]),
        ("6", "Налоги (НПД 6%)", [_fmt(t) for t in TAXES]),
        ("7", "Чистая прибыль", [_fmt(n) for n in NET_PROFIT]),
        (
            "8",
            "Рентабельность чистой прибыли, %",
            [str(round(NET_PROFIT[i] / REVENUES[i] * 100)) if REVENUES[i] else "0" for i in range(12)],
        ),
    ]

    t5 = _add_table(doc, 1 + len(table5_rows), 14)
    _style_cell(t5.rows[0].cells[0], "№", bold=True)
    _style_cell(t5.rows[0].cells[1], "Показатель", bold=True)
    for m in range(12):
        _style_cell(t5.rows[0].cells[m + 2], MONTH_LABELS[m], bold=True)
    for ri, (num, label, values) in enumerate(table5_rows, start=1):
        _style_cell(t5.rows[ri].cells[0], num)
        _style_cell(t5.rows[ri].cells[1], label)
        for m, val in enumerate(values):
            _style_cell(t5.rows[ri].cells[m + 2], val)
    doc.add_paragraph()

    # --- Таблица 6 ---
    _add_table_ref(doc, *FINANCE_TABLE_REFS[1])
    t6 = _add_table(doc, 1 + len(TABLE6_ROWS), 4)
    hdr6 = ["Показатель", "Ед. изм.", "Среднемесячно", "За год"]
    for j, h in enumerate(hdr6):
        _style_cell(t6.rows[0].cells[j], h, bold=True)
    for i, row in enumerate(TABLE6_ROWS, start=1):
        for j, val in enumerate(row):
            _style_cell(t6.rows[i].cells[j], val)
    doc.add_paragraph()

    # --- Таблица 7 ---
    _add_table_ref(doc, *FINANCE_TABLE_REFS[2])
    t7 = _add_table(doc, 3, 3)
    hdr7 = ["Источник", "Сумма, руб.", "Доля, %"]
    for j, h in enumerate(hdr7):
        _style_cell(t7.rows[0].cells[j], h, bold=True)
    _style_cell(t7.rows[1].cells[0], "Средства социального контракта")
    _style_cell(t7.rows[1].cells[1], "350 000")
    _style_cell(t7.rows[1].cells[2], "100")
    _style_cell(t7.rows[2].cells[0], "Собственные средства")
    _style_cell(t7.rows[2].cells[1], "0")
    _style_cell(t7.rows[2].cells[2], "0")
    doc.add_paragraph()

    # --- Таблица 8 ---
    _add_table_ref(doc, *FINANCE_TABLE_REFS[3])
    t8 = _add_table(doc, 1 + len(TABLE8_RISKS), 2)
    _style_cell(t8.rows[0].cells[0], "Риск", bold=True)
    _style_cell(t8.rows[0].cells[1], "Меры снижения", bold=True)
    for i, (risk, measure) in enumerate(TABLE8_RISKS, start=1):
        _style_cell(t8.rows[i].cells[0], risk)
        _style_cell(t8.rows[i].cells[1], measure)

    return doc


def verify_totals() -> list[str]:
    errors: list[str] = []
    if TABLE1_TOTAL != 350_000:
        errors.append(f"Таблица 1: ИТОГО {TABLE1_TOTAL}, ожидалось 350 000")
    if abs(TABLE3_TOTAL_REV - 300_000) > 500:
        errors.append(f"Таблица 3: выручка {TABLE3_TOTAL_REV}, ожидалось ~300 000")
    if REVENUES[-1] != 300_000:
        errors.append(f"Таблица 5: месяц 12 выручка {REVENUES[-1]}")
    if CAPEX_ROW[0] != 350_000:
        errors.append("Таблица 5: месяц 1 не содержит 350 000 в 4.2")
    if YEAR_REVENUE != 1_659_000:
        errors.append(f"Годовая выручка {YEAR_REVENUE}")
    if YEAR_NET != 1_030_510:
        errors.append(f"Чистая прибыль за год {YEAR_NET}, ожидалось 1 030 510")
    if TABLE1_SUMS["5"] != LIMIT_PROMO:
        errors.append(f"Продвижение: {TABLE1_SUMS['5']}, лимит {LIMIT_PROMO}")
    if TABLE1_SUMS["3"] != LIMIT_RENT:
        errors.append(f"Аренда: {TABLE1_SUMS['3']}, лимит {LIMIT_RENT}")
    if TABLE1_SUMS["4"] != LIMIT_SW:
        errors.append(f"ПО: {TABLE1_SUMS['4']}, лимит {LIMIT_SW}")
    return errors


def main() -> int:
    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)

    tpl = build_template()
    tpl.save(str(TEMPLATE_PATH))
    print(f"Шаблон: {TEMPLATE_PATH}")

    errs = verify_totals()
    if errs:
        print("ПРЕДУПРЕЖДЕНИЯ:", *errs, sep="\n  ")

    doc = fill_document(Document())
    for out in (OUTPUT_REGENERATED, OUTPUT_PATH):
        try:
            doc.save(str(out))
            print(f"Заполненный БП: {out}")
            break
        except PermissionError:
            continue
    else:
        print("Не удалось сохранить — закройте файлы docx в Word.")
        return 1
    return 1 if errs else 0


if __name__ == "__main__":
    raise SystemExit(main())
